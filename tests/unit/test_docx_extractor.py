"""Tests for the DOCX extractor."""

from pathlib import Path

import pytest

from opennarrator.exceptions import ExtractionError
from opennarrator.pipeline.extractor import DocxExtractor


@pytest.fixture
def docx_with_chapters(tmp_path: Path) -> Path:
    """Create a .docx with heading-based chapters."""
    from docx import Document

    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("This is the first chapter with enough content to be interesting.")
    doc.add_paragraph("It continues with more text in the same chapter.")
    doc.add_heading("Chapter 2", level=1)
    doc.add_paragraph("This is the second chapter with different content.")
    doc.add_paragraph("More text here to fill out the chapter.")
    doc.add_paragraph("Even more text for good measure.")
    doc.add_heading("Final Chapter", level=2)
    doc.add_paragraph("The last chapter wraps things up nicely.")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_flat(tmp_path: Path) -> Path:
    """Create a .docx with no headings — single chapter fallback."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(
        "This is a plain document with no heading styles. It should be treated as a single chapter."
    )
    doc.add_paragraph("The second paragraph continues the same chapter.")
    path = tmp_path / "flat.docx"
    doc.save(str(path))
    return path


class TestDocxExtractor:
    def test_extracts_heading_chapters(self, docx_with_chapters: Path) -> None:
        extractor = DocxExtractor()
        meta = extractor.extract(docx_with_chapters)
        assert len(meta.chapters) >= 2
        assert meta.chapters[0].title == "Chapter 1"
        assert meta.chapters[1].title == "Chapter 2"

    def test_fallback_single_chapter(self, docx_flat: Path) -> None:
        extractor = DocxExtractor()
        meta = extractor.extract(docx_flat)
        assert len(meta.chapters) == 1

    def test_extracts_metadata(self, docx_with_chapters: Path) -> None:
        extractor = DocxExtractor()
        meta = extractor.extract(docx_with_chapters)
        # python-docx sets these automatically
        assert meta.title
        assert meta.author

    def test_raises_on_missing_file(self) -> None:
        extractor = DocxExtractor()
        with pytest.raises(ExtractionError, match="File not found"):
            extractor.extract(Path("/nonexistent.docx"))

    def test_raises_on_wrong_format(self, tmp_path: Path) -> None:
        extractor = DocxExtractor()
        fake = tmp_path / "test.txt"
        fake.write_text("not a docx")
        with pytest.raises(ExtractionError, match="Expected .docx file"):
            extractor.extract(fake)

    def test_chapters_have_content(self, docx_with_chapters: Path) -> None:
        extractor = DocxExtractor()
        meta = extractor.extract(docx_with_chapters)
        for ch in meta.chapters:
            assert len(ch.text) > 0

    def test_chapter_indices_sequential(self, docx_with_chapters: Path) -> None:
        extractor = DocxExtractor()
        meta = extractor.extract(docx_with_chapters)
        for i, ch in enumerate(meta.chapters, start=1):
            assert ch.index == i
